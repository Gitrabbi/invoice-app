import os
import re
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
import tempfile
import shutil
from datetime import datetime
import subprocess
port = int(os.environ.get("PORT", 8501))  # Fallback to 8501


# ================== PATH CONFIGURATION ==================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "templates", "invoice_template.docx")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "invoice_pdfs")

# Create folders if they don't exist
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(os.path.join(BASE_DIR, "templates"), exist_ok=True)

# ================== HELPER FUNCTIONS ==================
def sanitize_filename(name):
    return re.sub(r'[\\/:*?"<>|]', '_', name)

def convert_docx_to_pdf(docx_path, pdf_path):
    """Cross-platform conversion using libreoffice"""
    try:
        # For Linux (Render)
        cmd = [
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', os.path.dirname(pdf_path), docx_path
        ]
        subprocess.run(cmd, check=True)
        return os.path.exists(pdf_path)
    except Exception as e:
        st.error(f"PDF conversion failed: {str(e)}")
        return False

def generate_pdf_from_template(template_path, row_data, output_folder, invoice_number):
    try:
        doc = Document(template_path)
        
        # Set default font size to 8pt for the entire document
        style = doc.styles['Normal']
        font = style.font
        font.size = Pt(8)
        
        # Format numeric values (keep empty strings for empty values)
        for key, value in row_data.items():
            if isinstance(value, (int, float)):
                row_data[key] = f"{value:.2f}" if pd.notna(value) else ""
        
        # Force Invoice Number and CURRENT Date
        current_date = datetime.now().strftime("%Y-%m-%d")
        row_data.update({
            "DATE": current_date,
            "INVOICE NUMBER": str(invoice_number)
        })
        
        # Add invoice number at the top
        if len(doc.paragraphs) > 0:
            p = doc.paragraphs[0]
            p.text = f"Invoice #: {invoice_number}\nDate: {current_date}\n" + p.text
            for run in p.runs:
                run.font.size = Pt(8)
        else:
            p = doc.add_paragraph(f"Invoice #: {invoice_number}\nDate: {current_date}")
            p.style.font.size = Pt(8)
        
        # Replace placeholders
    for paragraph in doc.paragraphs:
        for key, value in row_data.items():
        # Corrected placeholder formats
        for placeholder in [f"{{{key}.}}", f"{{{key}}}"]:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
                    run.font.size = Pt(8)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in row_data.items():
                       for placeholder in [f"{{key}.}}", f"{{{key}}}"]:
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
        
        # Save files
        customer_name = row_data.get("MARK", "Customer")
        contact_number = row_data.get("CONTACT NUMBER", "")
        invoice_total = row_data.get("TOTAL CHARGES_SUM", "0.00")
        
        pdf_name = f"Invoice_{invoice_number}_{sanitize_filename(customer_name)}_{sanitize_filename(contact_number)}_{sanitize_filename(invoice_total)}.pdf"
        temp_docx = os.path.join(output_folder, "temp.docx")
        pdf_path = os.path.join(output_folder, pdf_name)
        
        doc.save(temp_docx)
        if convert_docx_to_pdf(temp_docx, pdf_path):
            os.remove(temp_docx)
            update_notification_sheet(output_folder, pdf_name, customer_name, invoice_number, contact_number, invoice_total)
            return pdf_path
        return None
        
    except Exception as e:
        st.error(f"Template processing error: {str(e)}")
        raise

def update_notification_sheet(output_folder, pdf_name, customer_name, invoice_number, contact_number, invoice_total):
    sheet_path = os.path.join(output_folder, "Customer_Notification_Sheet.xlsx")
    file_url = os.path.join(output_folder, pdf_name)
    
    new_entry = pd.DataFrame([{
        "CUSTOMER": customer_name,
        "INVOICE URL": file_url,
        "INVOICE NO": invoice_number,
        "CONTACT NO": contact_number,
        "INVOICE TOTAL": invoice_total
    }])
    
    if os.path.exists(sheet_path):
        try:
            existing_data = pd.read_excel(sheet_path)
            updated_data = pd.concat([existing_data, new_entry], ignore_index=True)
        except:
            updated_data = new_entry
    else:
        updated_data = new_entry
    
    updated_data.to_excel(sheet_path, index=False)

def consolidate_rows(df):
    consolidated_data = []
    for customer_name, group in df.groupby("MARK"):
        total_qty = group["QTY"].sum(skipna=True)
        total_cbm = group["CBM"].sum(skipna=True)
        parking_charges = group["PARKING CHARGES"].dropna().iloc[0] if not group["PARKING CHARGES"].dropna().empty else 0
        
        if total_cbm < 0.05:
            calculated_charges = 10.00
        else:
            calculated_charges = (group["CBM"] * group["PER CHARGES"]).sum(skipna=True)
        
        total_charges = calculated_charges + parking_charges
        first_row = group.iloc[0]
        
        # Prepare multi-line values
        receipt_nos = [str(x) if pd.notna(x) else "" for x in group["RECEIPT NO."]]
        qtys = [f"{x:.2f}" if pd.notna(x) else "" for x in group["QTY"]]
        descriptions = [str(x) if pd.notna(x) else "" for x in group["DESCRIPTION"]]
        cbms = [f"{x:.2f}" if pd.notna(x) else "" for x in group["CBM"]]
        weights = [f"{x:.2f}" if pd.notna(x) else "" for x in group["WEIGHT(KG)"]]
        
        consolidated_data.append({
            "RECEIPT NO.": "\n".join(receipt_nos) if any(receipt_nos) else "",
            "QTY": "\n".join(qtys) if any(qtys) else "",
            "DESCRIPTION": "\n".join(descriptions) if any(descriptions) else "",
            "CBM": "\n".join(cbms) if any(cbms) else "",
            "WEIGHT(KG)": "\n".join(weights) if any(weights) else "",
            "PARKING CHARGES": f"{parking_charges:.2f}",
            "PER CHARGES": f"{first_row['PER CHARGES']:.2f}" if pd.notna(first_row['PER CHARGES']) else "",
            "TOTAL CHARGES": f"{total_charges:.2f}",
            "MARK": customer_name,
            "CONTACT NUMBER": str(first_row.get("CONTACT NUMBER", "")) if pd.notna(first_row.get("CONTACT NUMBER")) else "",
            "TOTAL QTY": f"{total_qty:.2f}",
            "TOTAL CBM": f"{total_cbm:.2f}",
            "TOTAL CHARGES_SUM": f"{total_charges:.2f}"
        })
    return consolidated_data

# ================== STREAMLIT UI ==================
st.title("PDF Invoice Generation Dashboard")

uploaded_file = st.file_uploader("Upload Excel Sheet", type=["xlsx", "xls"])
if uploaded_file is not None:
    try:
        df = pd.read_excel(uploaded_file)
        
        # Initialize missing columns
        for col in ["CARGO NUMBER", "TRACKING NUMBER", "TERMS", "PARKING CHARGES"]:
            if col not in df.columns:
                df[col] = "" if col != "PARKING CHARGES" else 0.0
        
        for col in ["Weight Rate", "PER CHARGES"]:
            if col not in df.columns:
                df[col] = 0.0

        # Session state management
        if 'edited_df' not in st.session_state:
            st.session_state.edited_df = df.copy()
        
        # Global updates section
        st.write("Global Updates:")
        global_per_charge = st.number_input("Default Per Charge Value", 
            value=float(df["PER CHARGES"].iloc[0]) if not df["PER CHARGES"].empty and pd.notna(df["PER CHARGES"].iloc[0]) else 0.0)
        
        if st.button("Apply Global Values"):
            st.session_state.edited_df = st.session_state.edited_df.copy()
            for customer in st.session_state.edited_df["MARK"].unique():
                st.session_state.edited_df.loc[st.session_state.edited_df["MARK"] == customer, "PER CHARGES"] = global_per_charge
            st.success("Global values applied!")

        # Data processing
        edited_df = st.session_state.edited_df.copy()
        edited_df["Weight CBM"] = edited_df.apply(
            lambda row: row["WEIGHT(KG)"] / row["Weight Rate"] if pd.notna(row["Weight Rate"]) and row["Weight Rate"] != 0 else 0,
            axis=1
        )
        edited_df["CBM"] = edited_df[["MEAS.(CBM)", "Weight CBM"]].max(axis=1)
        edited_df["Calculated Charges"] = edited_df["CBM"] * edited_df["PER CHARGES"]
        
        consolidated_df = pd.DataFrame(consolidate_rows(edited_df))
        st.write("Processed Data:", consolidated_df)
        
        # Invoice generation
        last_invoice_number = st.number_input("Last invoice number", min_value=0, value=0)
        
        if st.button("Generate All Invoices"):
            os.makedirs(OUTPUT_FOLDER, exist_ok=True)
            invoice_number = last_invoice_number + 1
            
            progress_bar = st.progress(0)
            for i, (_, row) in enumerate(consolidated_df.iterrows()):
                progress_bar.progress((i + 1) / len(consolidated_df))
                template_data = {k: v for k, v in row.items() if not k.startswith("TOTAL")}
                pdf_path = generate_pdf_from_template(TEMPLATE_PATH, template_data, OUTPUT_FOLDER, invoice_number + i)
            
            st.success(f"Generated {len(consolidated_df)} invoices in {OUTPUT_FOLDER}")
    
    except Exception as e:
        st.error(f"Error: {str(e)}")
